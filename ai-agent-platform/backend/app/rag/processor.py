"""
文档处理器

负责文档的解析、分块、预处理等功能。
"""

import re
import logging
from typing import List, Dict, Any, Optional, Tuple
from enum import Enum
import uuid
from datetime import datetime
from pathlib import Path

from .vectorstore import Document
from .embeddings import embedding_manager, TokenCounter

logger = logging.getLogger(__name__)


class ChunkStrategy(Enum):
    """分块策略"""
    FIXED_SIZE = "fixed_size"  # 固定大小分块
    SEMANTIC = "semantic"  # 语义分块
    SENTENCE = "sentence"  # 句子分块
    PARAGRAPH = "paragraph"  # 段落分块
    SLIDING_WINDOW = "sliding_window"  # 滑动窗口


class DocumentChunk:
    """文档块"""
    
    def __init__(self, content: str, metadata: Dict[str, Any] = None,
                 start_pos: int = 0, end_pos: int = 0):
        self.id = str(uuid.uuid4())
        self.content = content
        self.metadata = metadata or {}
        self.start_pos = start_pos
        self.end_pos = end_pos
        self.token_count = TokenCounter.count_tokens(content)
        self.created_at = datetime.now()
    
    def to_document(self, doc_id: str = None, embedding: List[float] = None) -> Document:
        """转换为Document对象"""
        return Document(
            id=doc_id or self.id,
            content=self.content,
            metadata=self.metadata,
            embedding=embedding
        )


class TextCleaner:
    """文本清理器"""
    
    @staticmethod
    def clean_text(text: str) -> str:
        """清理文本"""
        # 移除多余的空白字符
        text = re.sub(r'\s+', ' ', text)
        
        # 移除特殊字符（保留中文、英文、数字、基本标点）
        text = re.sub(r'[^\w\s\u4e00-\u9fff.,!?;:()[\]{}"\'-]', '', text)
        
        # 移除多余的换行符
        text = re.sub(r'\n+', '\n', text)
        
        # 去除首尾空白
        text = text.strip()
        
        return text
    
    @staticmethod
    def extract_metadata(text: str) -> Dict[str, Any]:
        """提取文本元数据"""
        metadata = {
            "length": len(text),
            "token_count": TokenCounter.count_tokens(text),
            "word_count": len(text.split()),
            "line_count": len(text.split('\n')),
            "has_chinese": bool(re.search(r'[\u4e00-\u9fff]', text)),
            "has_english": bool(re.search(r'[a-zA-Z]', text)),
            "has_numbers": bool(re.search(r'\d', text))
        }
        return metadata


class FixedSizeChunker:
    """固定大小分块器"""
    
    def __init__(self, chunk_size: int = 1000, overlap: int = 200):
        self.chunk_size = chunk_size
        self.overlap = overlap
    
    def chunk(self, text: str, metadata: Dict[str, Any] = None) -> List[DocumentChunk]:
        """分块文本"""
        chunks = []
        text_length = len(text)
        
        start = 0
        while start < text_length:
            end = min(start + self.chunk_size, text_length)
            
            # 尝试在句号处分割
            if end < text_length:
                last_period = text.rfind('。', start, end)
                last_dot = text.rfind('.', start, end)
                split_pos = max(last_period, last_dot)
                
                if split_pos > start + self.chunk_size // 2:
                    end = split_pos + 1
            
            chunk_text = text[start:end].strip()
            if chunk_text:
                chunk_metadata = metadata.copy() if metadata else {}
                chunk_metadata.update({
                    "chunk_index": len(chunks),
                    "start_pos": start,
                    "end_pos": end,
                    "chunk_strategy": "fixed_size"
                })
                
                chunk = DocumentChunk(
                    content=chunk_text,
                    metadata=chunk_metadata,
                    start_pos=start,
                    end_pos=end
                )
                chunks.append(chunk)
            
            # 计算下一个起始位置（考虑重叠）
            start = max(start + self.chunk_size - self.overlap, end)
        
        return chunks


class SemanticChunker:
    """语义分块器"""
    
    def __init__(self, max_chunk_size: int = 1000, similarity_threshold: float = 0.7):
        self.max_chunk_size = max_chunk_size
        self.similarity_threshold = similarity_threshold
    
    def chunk(self, text: str, metadata: Dict[str, Any] = None) -> List[DocumentChunk]:
        """语义分块"""
        # 先按段落分割
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        if not paragraphs:
            return []
        
        chunks = []
        current_chunk = ""
        current_start = 0
        
        for i, paragraph in enumerate(paragraphs):
            # 检查添加这个段落后是否超过大小限制
            test_chunk = current_chunk + "\n\n" + paragraph if current_chunk else paragraph
            
            if TokenCounter.count_tokens(test_chunk) <= self.max_chunk_size:
                current_chunk = test_chunk
            else:
                # 保存当前块
                if current_chunk:
                    chunk_metadata = metadata.copy() if metadata else {}
                    chunk_metadata.update({
                        "chunk_index": len(chunks),
                        "chunk_strategy": "semantic",
                        "paragraph_count": current_chunk.count('\n\n') + 1
                    })
                    
                    chunk = DocumentChunk(
                        content=current_chunk,
                        metadata=chunk_metadata,
                        start_pos=current_start,
                        end_pos=current_start + len(current_chunk)
                    )
                    chunks.append(chunk)
                
                # 开始新块
                current_chunk = paragraph
                current_start = text.find(paragraph, current_start)
        
        # 添加最后一个块
        if current_chunk:
            chunk_metadata = metadata.copy() if metadata else {}
            chunk_metadata.update({
                "chunk_index": len(chunks),
                "chunk_strategy": "semantic",
                "paragraph_count": current_chunk.count('\n\n') + 1
            })
            
            chunk = DocumentChunk(
                content=current_chunk,
                metadata=chunk_metadata,
                start_pos=current_start,
                end_pos=current_start + len(current_chunk)
            )
            chunks.append(chunk)
        
        return chunks


class SentenceChunker:
    """句子分块器"""
    
    def __init__(self, sentences_per_chunk: int = 5, overlap_sentences: int = 1):
        self.sentences_per_chunk = sentences_per_chunk
        self.overlap_sentences = overlap_sentences
    
    def chunk(self, text: str, metadata: Dict[str, Any] = None) -> List[DocumentChunk]:
        """按句子分块"""
        # 分割句子
        sentences = self._split_sentences(text)
        
        if not sentences:
            return []
        
        chunks = []
        start_idx = 0
        
        while start_idx < len(sentences):
            end_idx = min(start_idx + self.sentences_per_chunk, len(sentences))
            
            chunk_sentences = sentences[start_idx:end_idx]
            chunk_text = ' '.join(chunk_sentences)
            
            chunk_metadata = metadata.copy() if metadata else {}
            chunk_metadata.update({
                "chunk_index": len(chunks),
                "chunk_strategy": "sentence",
                "sentence_count": len(chunk_sentences),
                "sentence_start": start_idx,
                "sentence_end": end_idx - 1
            })
            
            chunk = DocumentChunk(
                content=chunk_text,
                metadata=chunk_metadata
            )
            chunks.append(chunk)
            
            # 计算下一个起始位置（考虑重叠）
            start_idx = max(start_idx + self.sentences_per_chunk - self.overlap_sentences, end_idx)
        
        return chunks
    
    def _split_sentences(self, text: str) -> List[str]:
        """分割句子"""
        # 简单的句子分割（可以使用更复杂的NLP库）
        sentences = []
        
        # 按中文句号分割
        parts = re.split(r'[。！？]', text)
        for part in parts:
            part = part.strip()
            if part:
                # 再按英文句号分割
                sub_parts = re.split(r'[.!?]', part)
                for sub_part in sub_parts:
                    sub_part = sub_part.strip()
                    if sub_part and len(sub_part) > 10:  # 过滤太短的句子
                        sentences.append(sub_part)
        
        return sentences


class DocumentProcessor:
    """文档处理器"""
    
    def __init__(self):
        self.text_cleaner = TextCleaner()
        self.chunkers = {
            ChunkStrategy.FIXED_SIZE: FixedSizeChunker(),
            ChunkStrategy.SEMANTIC: SemanticChunker(),
            ChunkStrategy.SENTENCE: SentenceChunker()
        }
    
    def process_text(self, text: str, strategy: ChunkStrategy = ChunkStrategy.FIXED_SIZE,
                    metadata: Dict[str, Any] = None, clean_text: bool = True) -> List[DocumentChunk]:
        """处理文本"""
        try:
            # 清理文本
            if clean_text:
                text = self.text_cleaner.clean_text(text)
            
            # 提取元数据
            text_metadata = self.text_cleaner.extract_metadata(text)
            if metadata:
                text_metadata.update(metadata)
            
            # 分块
            chunker = self.chunkers.get(strategy)
            if not chunker:
                raise ValueError(f"不支持的分块策略: {strategy}")
            
            chunks = chunker.chunk(text, text_metadata)
            
            logger.info(f"文本处理完成: {len(chunks)} 个块, 策略: {strategy.value}")
            return chunks
            
        except Exception as e:
            logger.error(f"文本处理失败: {e}")
            raise
    
    def process_file(self, file_path: str, strategy: ChunkStrategy = ChunkStrategy.FIXED_SIZE,
                    metadata: Dict[str, Any] = None) -> List[DocumentChunk]:
        """处理文件"""
        try:
            # 读取文件内容
            text = self._read_file(file_path)
            
            # 添加文件元数据
            file_metadata = {
                "file_path": file_path,
                "file_name": Path(file_path).name,
                "file_extension": Path(file_path).suffix,
                "processed_at": datetime.now().isoformat()
            }
            
            if metadata:
                file_metadata.update(metadata)
            
            return self.process_text(text, strategy, file_metadata)
            
        except Exception as e:
            logger.error(f"文件处理失败: {e}")
            raise
    
    def _read_file(self, file_path: str) -> str:
        """读取文件内容"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        extension = file_path.suffix.lower()
        
        if extension == '.txt':
            return self._read_txt(file_path)
        elif extension == '.pdf':
            return self._read_pdf(file_path)
        elif extension in ['.docx', '.doc']:
            return self._read_docx(file_path)
        else:
            # 尝试作为文本文件读取
            try:
                return self._read_txt(file_path)
            except:
                raise ValueError(f"不支持的文件格式: {extension}")
    
    def _read_txt(self, file_path: Path) -> str:
        """读取文本文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # 尝试其他编码
            for encoding in ['gbk', 'gb2312', 'latin-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            raise ValueError("无法解码文件")
    
    def _read_pdf(self, file_path: Path) -> str:
        """读取PDF文件"""
        try:
            import PyPDF2
            text = ""
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            return text
        except ImportError:
            raise ImportError("需要安装PyPDF2: pip install PyPDF2")
        except Exception as e:
            logger.error(f"PDF读取失败: {e}")
            raise
    
    def _read_docx(self, file_path: Path) -> str:
        """读取DOCX文件"""
        try:
            from docx import Document
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except ImportError:
            raise ImportError("需要安装python-docx: pip install python-docx")
        except Exception as e:
            logger.error(f"DOCX读取失败: {e}")
            raise
    
    async def process_and_embed(self, text: str, strategy: ChunkStrategy = ChunkStrategy.FIXED_SIZE,
                               metadata: Dict[str, Any] = None, model_name: str = None) -> List[Document]:
        """处理文本并生成嵌入"""
        # 处理文本
        chunks = self.process_text(text, strategy, metadata)
        
        # 生成嵌入
        texts = [chunk.content for chunk in chunks]
        embeddings = await embedding_manager.embed_texts(texts, model_name)
        
        # 创建Document对象
        documents = []
        for chunk, embedding in zip(chunks, embeddings):
            doc = chunk.to_document(embedding=embedding)
            documents.append(doc)
        
        return documents
    
    def configure_chunker(self, strategy: ChunkStrategy, **kwargs):
        """配置分块器"""
        if strategy == ChunkStrategy.FIXED_SIZE:
            self.chunkers[strategy] = FixedSizeChunker(**kwargs)
        elif strategy == ChunkStrategy.SEMANTIC:
            self.chunkers[strategy] = SemanticChunker(**kwargs)
        elif strategy == ChunkStrategy.SENTENCE:
            self.chunkers[strategy] = SentenceChunker(**kwargs)


# 全局文档处理器
document_processor = DocumentProcessor()
