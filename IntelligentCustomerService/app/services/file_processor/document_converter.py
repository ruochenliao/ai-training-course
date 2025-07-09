"""
文档转换器
支持多种文档格式转换为Markdown
"""
import asyncio
import csv
import json
import logging
import os
import time
from pathlib import Path
from typing import Dict, List, Any, Optional

try:
    # PDF处理
    from marker import convert_single_pdf
    MARKER_AVAILABLE = True
except ImportError:
    MARKER_AVAILABLE = False

try:
    # Office文档处理
    from docx import Document as DocxDocument
    from openpyxl import load_workbook
    from pptx import Presentation
    OFFICE_AVAILABLE = True
except ImportError:
    OFFICE_AVAILABLE = False

try:
    # HTML处理
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False

try:
    # 数据处理
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

from .processor_config import ProcessorConfig, FileType, ConversionMethod

logger = logging.getLogger(__name__)


class DocumentConverter:
    """
    文档转换器
    支持多种文档格式转换为Markdown
    """
    
    def __init__(self, config: ProcessorConfig):
        """
        初始化文档转换器
        
        Args:
            config: 处理器配置
        """
        self.config = config
        self.initialized = False
        
        # 统计信息
        self.stats = {
            "total_conversions": 0,
            "successful_conversions": 0,
            "failed_conversions": 0,
            "conversion_times": {},
            "error_counts": {}
        }
        
        logger.info("文档转换器初始化完成")
    
    async def initialize(self) -> None:
        """初始化转换器"""
        try:
            # 检查依赖库
            self._check_dependencies()
            
            self.initialized = True
            logger.info("文档转换器初始化成功")
            
        except Exception as e:
            logger.error(f"文档转换器初始化失败: {e}")
            raise
    
    def _check_dependencies(self) -> None:
        """检查依赖库"""
        missing_deps = []
        
        if not MARKER_AVAILABLE:
            missing_deps.append("marker-pdf (pip install marker-pdf)")
        
        if not OFFICE_AVAILABLE:
            missing_deps.append("office libraries (pip install python-docx openpyxl python-pptx)")
        
        if not BS4_AVAILABLE:
            missing_deps.append("beautifulsoup4 (pip install beautifulsoup4)")
        
        if not PANDAS_AVAILABLE:
            missing_deps.append("pandas (pip install pandas)")
        
        if missing_deps:
            logger.warning(f"缺少可选依赖库: {', '.join(missing_deps)}")
    
    async def convert_to_markdown(
        self,
        file_path: str,
        file_type: FileType,
        conversion_method: ConversionMethod,
        metadata: Dict[str, Any] = None
    ) -> Dict[str, Any]:
        """
        转换文档为Markdown
        
        Args:
            file_path: 文件路径
            file_type: 文件类型
            conversion_method: 转换方法
            metadata: 元数据
            
        Returns:
            转换结果
        """
        start_time = time.time()
        
        try:
            if not self.initialized:
                raise RuntimeError("文档转换器未初始化")
            
            self.stats["total_conversions"] += 1
            
            # 根据转换方法选择处理函数
            if conversion_method == ConversionMethod.MARKER_PDF:
                result = await self._convert_pdf_with_marker(file_path, metadata)
            elif conversion_method == ConversionMethod.PYTHON_DOCX:
                result = await self._convert_docx(file_path, metadata)
            elif conversion_method == ConversionMethod.OPENPYXL:
                result = await self._convert_xlsx(file_path, metadata)
            elif conversion_method == ConversionMethod.PYTHON_PPTX:
                result = await self._convert_pptx(file_path, metadata)
            elif conversion_method == ConversionMethod.DIRECT_READ:
                result = await self._convert_text_file(file_path, file_type, metadata)
            elif conversion_method == ConversionMethod.PANDAS:
                result = await self._convert_csv(file_path, metadata)
            elif conversion_method == ConversionMethod.BEAUTIFULSOUP:
                result = await self._convert_html(file_path, metadata)
            else:
                raise ValueError(f"不支持的转换方法: {conversion_method.value}")
            
            # 记录统计信息
            processing_time = time.time() - start_time
            self._update_stats(conversion_method, processing_time, True)
            
            result["processing_time"] = processing_time
            result["conversion_method"] = conversion_method.value
            
            logger.info(f"文档转换成功: {file_path}, 方法: {conversion_method.value}, 耗时: {processing_time:.2f}秒")
            
            return result
            
        except Exception as e:
            # 记录错误统计
            processing_time = time.time() - start_time
            self._update_stats(conversion_method, processing_time, False)
            
            logger.error(f"文档转换失败: {file_path}, 方法: {conversion_method.value}, 错误: {e}")
            
            return {
                "success": False,
                "error": str(e),
                "content": "",
                "metadata": {},
                "processing_time": processing_time,
                "conversion_method": conversion_method.value
            }
    
    async def _convert_pdf_with_marker(self, file_path: str, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """使用marker-pdf转换PDF"""
        if not MARKER_AVAILABLE:
            raise ImportError("marker-pdf库不可用")
        
        try:
            # 在线程池中执行Marker转换
            loop = asyncio.get_event_loop()
            result = await loop.run_in_executor(
                None,
                self._convert_pdf_sync,
                file_path
            )
            
            return {
                "success": True,
                "content": result["markdown"],
                "metadata": {
                    "pages": result.get("pages", 0),
                    "images": result.get("images", []),
                    "tables": result.get("tables", []),
                    "marker_metadata": result.get("metadata", {})
                }
            }
            
        except Exception as e:
            logger.error(f"Marker PDF转换失败: {e}")
            raise
    
    def _convert_pdf_sync(self, file_path: str) -> Dict[str, Any]:
        """同步执行Marker PDF转换"""
        try:
            # 使用Marker转换PDF
            full_text, images, out_meta = convert_single_pdf(
                file_path,
                max_pages=self.config.marker_config.get("max_pages"),
                langs=self.config.marker_config.get("langs", ["zh", "en"]),
                batch_multiplier=self.config.marker_config.get("batch_multiplier", 2)
            )
            
            return {
                "markdown": full_text,
                "images": images,
                "metadata": out_meta,
                "pages": out_meta.get("page_count", 0)
            }
            
        except Exception as e:
            logger.error(f"Marker同步转换失败: {e}")
            raise
    
    async def _convert_docx(self, file_path: str, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """转换DOCX文档"""
        if not OFFICE_AVAILABLE:
            raise ImportError("python-docx库不可用")
        
        try:
            # 在线程池中执行转换
            loop = asyncio.get_event_loop()
            result = await loop.run_in_executor(
                None,
                self._convert_docx_sync,
                file_path
            )
            
            return {
                "success": True,
                "content": result["markdown"],
                "metadata": result["metadata"]
            }
            
        except Exception as e:
            logger.error(f"DOCX转换失败: {e}")
            raise
    
    def _convert_docx_sync(self, file_path: str) -> Dict[str, Any]:
        """同步转换DOCX文档"""
        try:
            doc = DocxDocument(file_path)
            
            markdown_parts = []
            tables_count = 0
            images_count = 0
            
            # 处理段落
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    # 检查是否是标题
                    if paragraph.style.name.startswith('Heading'):
                        level = int(paragraph.style.name.split()[-1])
                        markdown_parts.append(f"{'#' * level} {text}")
                    else:
                        markdown_parts.append(text)
                    markdown_parts.append("")  # 添加空行
            
            # 处理表格
            for table in doc.tables:
                tables_count += 1
                markdown_parts.append(f"\n## 表格 {tables_count}\n")
                
                # 表格标题行
                if table.rows:
                    header_cells = [cell.text.strip() for cell in table.rows[0].cells]
                    markdown_parts.append("| " + " | ".join(header_cells) + " |")
                    markdown_parts.append("| " + " | ".join(["---"] * len(header_cells)) + " |")
                    
                    # 表格数据行
                    for row in table.rows[1:]:
                        data_cells = [cell.text.strip() for cell in row.cells]
                        markdown_parts.append("| " + " | ".join(data_cells) + " |")
                
                markdown_parts.append("")
            
            markdown_content = "\n".join(markdown_parts)
            
            return {
                "markdown": markdown_content,
                "metadata": {
                    "paragraphs": len(doc.paragraphs),
                    "tables": tables_count,
                    "images": images_count,
                    "word_count": len(markdown_content.split())
                }
            }
            
        except Exception as e:
            logger.error(f"DOCX同步转换失败: {e}")
            raise
    
    async def _convert_xlsx(self, file_path: str, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """转换XLSX文档"""
        if not OFFICE_AVAILABLE:
            raise ImportError("openpyxl库不可用")
        
        try:
            # 在线程池中执行转换
            loop = asyncio.get_event_loop()
            result = await loop.run_in_executor(
                None,
                self._convert_xlsx_sync,
                file_path
            )
            
            return {
                "success": True,
                "content": result["markdown"],
                "metadata": result["metadata"]
            }
            
        except Exception as e:
            logger.error(f"XLSX转换失败: {e}")
            raise
    
    def _convert_xlsx_sync(self, file_path: str) -> Dict[str, Any]:
        """同步转换XLSX文档"""
        try:
            workbook = load_workbook(file_path, data_only=True)
            
            markdown_parts = []
            total_rows = 0
            total_cols = 0
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                
                markdown_parts.append(f"# 工作表: {sheet_name}\n")
                
                # 获取有数据的范围
                if sheet.max_row > 0 and sheet.max_column > 0:
                    # 读取数据
                    data = []
                    for row in sheet.iter_rows(values_only=True):
                        # 过滤空行
                        if any(cell is not None for cell in row):
                            data.append([str(cell) if cell is not None else "" for cell in row])
                    
                    if data:
                        # 创建Markdown表格
                        if len(data) > 0:
                            # 表头
                            header = data[0]
                            markdown_parts.append("| " + " | ".join(header) + " |")
                            markdown_parts.append("| " + " | ".join(["---"] * len(header)) + " |")
                            
                            # 数据行
                            for row in data[1:]:
                                # 确保行长度与表头一致
                                while len(row) < len(header):
                                    row.append("")
                                markdown_parts.append("| " + " | ".join(row[:len(header)]) + " |")
                        
                        total_rows += len(data)
                        total_cols = max(total_cols, len(data[0]) if data else 0)
                
                markdown_parts.append("")
            
            markdown_content = "\n".join(markdown_parts)
            
            return {
                "markdown": markdown_content,
                "metadata": {
                    "sheets": len(workbook.sheetnames),
                    "total_rows": total_rows,
                    "total_columns": total_cols,
                    "sheet_names": workbook.sheetnames
                }
            }
            
        except Exception as e:
            logger.error(f"XLSX同步转换失败: {e}")
            raise

    async def _convert_html(self, file_path: str, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """转换HTML文件"""
        if not BS4_AVAILABLE:
            # 降级到直接读取
            return await self._convert_text_file(file_path, FileType.HTML, metadata)

        try:
            # 在线程池中执行转换
            loop = asyncio.get_event_loop()
            result = await loop.run_in_executor(
                None,
                self._convert_html_sync,
                file_path
            )

            return {
                "success": True,
                "content": result["markdown"],
                "metadata": result["metadata"]
            }

        except Exception as e:
            logger.error(f"HTML转换失败: {e}")
            raise

    def _convert_html_sync(self, file_path: str) -> Dict[str, Any]:
        """同步转换HTML文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()

            soup = BeautifulSoup(html_content, 'html.parser')

            # 提取文本内容并转换为Markdown
            markdown_parts = []

            # 处理标题
            for i in range(1, 7):
                for heading in soup.find_all(f'h{i}'):
                    text = heading.get_text().strip()
                    if text:
                        markdown_parts.append(f"{'#' * i} {text}")
                        markdown_parts.append("")

            # 处理段落
            for p in soup.find_all('p'):
                text = p.get_text().strip()
                if text:
                    markdown_parts.append(text)
                    markdown_parts.append("")

            # 处理列表
            for ul in soup.find_all('ul'):
                for li in ul.find_all('li'):
                    text = li.get_text().strip()
                    if text:
                        markdown_parts.append(f"- {text}")
                markdown_parts.append("")

            for ol in soup.find_all('ol'):
                for i, li in enumerate(ol.find_all('li'), 1):
                    text = li.get_text().strip()
                    if text:
                        markdown_parts.append(f"{i}. {text}")
                markdown_parts.append("")

            # 处理表格
            for table in soup.find_all('table'):
                rows = table.find_all('tr')
                if rows:
                    # 表头
                    header_row = rows[0]
                    headers = [th.get_text().strip() for th in header_row.find_all(['th', 'td'])]
                    if headers:
                        markdown_parts.append("| " + " | ".join(headers) + " |")
                        markdown_parts.append("| " + " | ".join(["---"] * len(headers)) + " |")

                        # 数据行
                        for row in rows[1:]:
                            cells = [td.get_text().strip() for td in row.find_all(['td', 'th'])]
                            while len(cells) < len(headers):
                                cells.append("")
                            markdown_parts.append("| " + " | ".join(cells[:len(headers)]) + " |")

                        markdown_parts.append("")

            markdown_content = "\n".join(markdown_parts)

            return {
                "markdown": markdown_content,
                "metadata": {
                    "title": soup.title.string if soup.title else "",
                    "headings": len(soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])),
                    "paragraphs": len(soup.find_all('p')),
                    "tables": len(soup.find_all('table')),
                    "links": len(soup.find_all('a')),
                    "images": len(soup.find_all('img'))
                }
            }

        except Exception as e:
            logger.error(f"HTML同步转换失败: {e}")
            raise

    def _update_stats(self, method: ConversionMethod, processing_time: float, success: bool) -> None:
        """
        更新统计信息

        Args:
            method: 转换方法
            processing_time: 处理时间
            success: 是否成功
        """
        method_name = method.value

        if success:
            self.stats["successful_conversions"] += 1
        else:
            self.stats["failed_conversions"] += 1
            self.stats["error_counts"][method_name] = self.stats["error_counts"].get(method_name, 0) + 1

        # 更新处理时间统计
        if method_name not in self.stats["conversion_times"]:
            self.stats["conversion_times"][method_name] = []

        self.stats["conversion_times"][method_name].append(processing_time)

        # 保持最近100次记录
        if len(self.stats["conversion_times"][method_name]) > 100:
            self.stats["conversion_times"][method_name] = self.stats["conversion_times"][method_name][-100:]

    def get_stats(self) -> Dict[str, Any]:
        """
        获取转换器统计信息

        Returns:
            统计信息
        """
        # 计算平均处理时间
        avg_times = {}
        for method, times in self.stats["conversion_times"].items():
            if times:
                avg_times[method] = sum(times) / len(times)

        return {
            **self.stats,
            "average_conversion_times": avg_times,
            "success_rate": (
                self.stats["successful_conversions"] / self.stats["total_conversions"]
                if self.stats["total_conversions"] > 0 else 0
            ),
            "dependencies": {
                "marker_available": MARKER_AVAILABLE,
                "office_available": OFFICE_AVAILABLE,
                "bs4_available": BS4_AVAILABLE,
                "pandas_available": PANDAS_AVAILABLE
            }
        }

    async def close(self) -> None:
        """关闭转换器，释放资源"""
        try:
            self.initialized = False
            logger.info("文档转换器已关闭")
        except Exception as e:
            logger.error(f"关闭文档转换器失败: {e}")
